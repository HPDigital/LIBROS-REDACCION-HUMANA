"""
LIBROS REDACCION HUMANA
"""

#!/usr/bin/env python
# coding: utf-8

# In[ ]:





# In[ ]:





# In[2]:


import os
from docx import Document
from dotenv import load_dotenv
from langchain.prompts import PromptTemplate
from langchain_openai.chat_models import ChatOpenAI
from PyPDF2 import PdfReader
import nltk
from nltk.tokenize import sent_tokenize
from langchain_openai import OpenAIEmbeddings  # Importación actualizada
from langchain.vectorstores import FAISS

# Descargar recursos de NLTK (si es la primera vez)
# nltk.download('punkt')

# =============================================================================
# CONFIGURACIÓN Y CLIENTES
# =============================================================================

load_dotenv()
API_KEY = os.getenv("OPENAI_API_KEY")
if not API_KEY:
    raise ValueError("La clave de la API de OpenAI no se encontró en las variables de entorno.")

# Inicializar cliente LLM
llm = ChatOpenAI(api_key=API_KEY, model="o3-mini-2025-01-31", temperature=1)

# Inicializar embeddings para vectorización (actualizado)
embeddings = OpenAIEmbeddings(api_key=API_KEY)

# =============================================================================
# TEMPLATES DE PROMPT
# =============================================================================

# Prompt para obtener el estilo a partir de un ejemplo en PDF
PROMPT_ESTILO = PromptTemplate(
    input_variables=["capitulo", "pestania", "grupo", "funcion", "texto_estilo", "contenido"],
    template="""Ajusta tu forma de redacción para cumplir con los siguientes criterios, asegurando un texto que se asemeje a una respuesta completamente humana y alejada de los patrones de escritura típicos de inteligencia artificial.

Restricción de Vocabulario:
Evita las 20 palabras más comunes en tus respuestas estándar, ya que buscamos mayor riqueza y variedad léxica. No uses: ChatGPT, modelo, respuestas, lenguaje, desarrollado, OpenAI, objetivo, proporcionar, claras, concisas, útiles, análisis, datos, programación, herramienta, versátil, capacidad, temas, interacción y estructura. En su lugar, emplea sinónimos o reformula las ideas sin comprometer su significado ni fluidez.

Restricción de Estructuras Gramaticales:
No utilices las siguinetes construcciones gramaticales más frecuentes en tu redacción típica:
Oraciones compuestas coordinadas con conectores predecibles.
Oraciones subordinadas explicativas tradicionales.
Oraciones impersonales con se.
Oraciones pasivas perifrásticas.

Sustitúyelas por alternativas más dinámicas y naturales, como:
Expresiones conversacionales con pausas estratégicas y énfasis natural.
Oraciones largas con múltiples proposiciones y uso variado de puntuación.
Enunciados exclamativos o interjecciones para mayor espontaneidad.
Uso de la segunda persona para generar cercanía con el lector.
Cada frase debe tener al menos tres proposiciones interconectadas y un mínimo de 50 palabras para garantizar profundidad y desarrollo adecuado.

Estilo de Redacción:
Imita la fluidez, la coherencia y la sofisticación de un texto académico extraído de un documento formal en español. Para ello, ten en cuenta las siguientes directrices:

Extensión y Complejidad de las Frases:
Redacta oraciones extensas que integren varias ideas de manera armónica, utilizando conectores fluidos en español.

Estructura Cohesiva y Fragmentación Adecuada:
Organiza el contenido en un mínimo de dos y un máximo de cinco oraciones interconectadas, asegurando una redacción natural y evitando frases inconexas.

Claridad y Consistencia:
Cada oración debe mantener un balance entre precisión y fluidez, evitando repeticiones innecesarias o explicaciones redundantes.

Uso Estratégico de Conectores:
Minimiza los puntos aislados y favorece una conexión fluida de ideas mediante conjunciones variadas, sin caer en estructuras mecánicas.

Diversificación Sintáctica:
Utiliza una combinación de oraciones subordinadas, construcciones con participios y elementos discursivos que den variedad y riqueza al texto.

Tono Académico pero Accesible:
Mantén un estilo formal y estructurado, pero sin tecnicismos excesivos que dificulten la comprensión. Piensa en la redacción de un libro universitario.

Extensión de las Frases:
Cada frase debe contener más de 50 palabras, integrando al menos tres proposiciones, asegurando un flujo de información profundo y detallado.

Instrucciones Finales:
Redacta el siguiente contenido aplicando todas las reglas anteriores, asegurando que el resultado sea indistinguible de un texto humano de alto nivel académico. Evita frases cortas, patrones predecibles o un tono excesivamente robótico. Usa la referencia proporcionada a continuación como base para emular el estilo lingüístico y discursivo.

Ejemplos de redaccion: Utiliza de ejemplo de redaccion las frases siguientes {texto_estilo}

Referencia de Estilo:
{texto_estilo}

Texto de Entrada:
{contenido}

Salida esperada:
Un texto extenso, fluido, formal y académico, que siga las directrices mencionadas, con variedad sintáctica y vocabulario enriquecido"""
)

# PROMPTS PARA CADA PARTE DEL RESUMEN

PROMPT_PROPOSITO = PromptTemplate(
    input_variables=["capitulo", "pestania", "grupo", "funcion"],
    template="""Proporciona una expliacion clara y detallada del uso de la función de Excel '{funcion}' del grupo '{grupo}' en la pestaña '{pestania}', en el contexto del capítulo '{capitulo}'. 
Incluye una explicación paso a paso del proceso, el resultado esperado y, de ser pertinente, menciona errores comunes junto con sus soluciones.
Ajusta tu forma de redacción para cumplir con los siguientes criterios, asegurando un texto que se asemeje a una respuesta completamente humana y alejada de los patrones de escritura típicos de inteligencia artificial.

Restricción de Vocabulario:
Evita las 20 palabras más comunes en tus respuestas estándar, ya que buscamos mayor riqueza y variedad léxica. No uses: ChatGPT, modelo, respuestas, lenguaje, desarrollado, OpenAI, objetivo, proporcionar, claras, concisas, útiles, análisis, datos, programación, herramienta, versátil, capacidad, temas, interacción y estructura. En su lugar, emplea sinónimos o reformula las ideas sin comprometer su significado ni fluidez.

Restricción de Estructuras Gramaticales:
No utilices las siguientes construcciones gramaticales más frecuentes en tu redacción típica:

Oraciones compuestas coordinadas con conectores predecibles.
Oraciones subordinadas explicativas tradicionales.
Oraciones impersonales con se.
Oraciones pasivas perifrásticas

Sustitúyelas por alternativas más dinámicas y naturales, como:
Expresiones conversacionales con pausas estratégicas y énfasis natural.
Oraciones largas con múltiples proposiciones y uso variado de puntuación.
Enunciados exclamativos o interjecciones para mayor espontaneidad.
Uso de la segunda persona para generar cercanía con el lector.
Cada frase debe tener al menos tres proposiciones interconectadas y un mínimo de 50 palabras para garantizar profundidad y desarrollo adecuado.

Estilo de Redacción:
Imita la fluidez, la coherencia y la sofisticación de un texto académico extraído de un documento formal en español. Para ello, ten en cuenta las siguientes directrices:

Extensión y Complejidad de las Frases:
Redacta oraciones extensas que integren varias ideas de manera armónica, utilizando conectores fluidos en español.

Estructura Cohesiva y Fragmentación Adecuada:
Organiza el contenido en un mínimo de dos y un máximo de cinco oraciones interconectadas, asegurando una redacción natural y evitando frases inconexas.

Claridad y Consistencia:
Cada oración debe mantener un balance entre precisión y fluidez, evitando repeticiones innecesarias o explicaciones redundantes.

Uso Estratégico de Conectores:
Minimiza los puntos aislados y favorece una conexión fluida de ideas mediante conjunciones variadas, sin caer en estructuras mecánicas.

Diversificación Sintáctica:
Utiliza una combinación de oraciones subordinadas, construcciones con participios y elementos discursivos que den variedad y riqueza al texto.

Tono Académico pero Accesible:
Mantén un estilo formal y estructurado, pero sin tecnicismos excesivos que dificulten la comprensión. Piensa en la redacción de un libro universitario.

Extensión de las Frases:
Cada frase debe contener más de 50 palabras, integrando al menos tres proposiciones, asegurando un flujo de información profundo y detallado.

Instrucciones Finales:
Redacta el siguiente contenido aplicando todas las reglas anteriores, asegurando que el resultado sea indistinguible de un texto humano de alto nivel académico. Evita frases cortas, patrones predecibles o un tono excesivamente robótico. Usa la referencia proporcionada como base para emular el estilo lingüístico y discursivo.



Salida esperada:
Un texto extenso, fluido, formal y académico, que siga las directrices mencionadas, con variedad sintáctica y vocabulario enriquecido
La respuesta debe ser clara, profesional y limitada a un máximo de 300 tokens.
La respueste debe tener frases de mas de 50 palbras que contengan varias oraciones.
Las frases deben estar compuestas por minimo tres oraciones.
Respuesta:"""
)

PROMPT_ENCAMINAMIENTO = PromptTemplate(
    input_variables=["capitulo", "pestania", "grupo", "funcion"],
    template="""Proporciona una explicación tecnica y detallado del uso de la función de Excel '{funcion}' del grupo '{grupo}' en la pestaña '{pestania}', en el contexto del capítulo '{capitulo}'. 
Incluye una explicación paso a paso del proceso, el resultado esperado y, de ser pertinente, menciona errores comunes junto con sus soluciones.
Ajusta tu forma de redacción para cumplir con los siguientes criterios, asegurando un texto que se asemeje a una respuesta completamente humana y alejada de los patrones de escritura típicos de inteligencia artificial.

Restricción de Vocabulario:
Evita las 20 palabras más comunes en tus respuestas estándar, ya que buscamos mayor riqueza y variedad léxica. No uses: ChatGPT, modelo, respuestas, lenguaje, desarrollado, OpenAI, objetivo, proporcionar, claras, concisas, útiles, análisis, datos, programación, herramienta, versátil, capacidad, temas, interacción y estructura. En su lugar, emplea sinónimos o reformula las ideas sin comprometer su significado ni fluidez.

Restricción de Estructuras Gramaticales:
No utilices las siguientes construcciones gramaticales más frecuentes en tu redacción típica:

Oraciones compuestas coordinadas con conectores predecibles.
Oraciones subordinadas explicativas tradicionales.
Oraciones impersonales con se.
Oraciones pasivas perifrásticas

Sustitúyelas por alternativas más dinámicas y naturales, como:
Expresiones conversacionales con pausas estratégicas y énfasis natural.
Oraciones largas con múltiples proposiciones y uso variado de puntuación.
Enunciados exclamativos o interjecciones para mayor espontaneidad.
Uso de la segunda persona para generar cercanía con el lector.
Cada frase debe tener al menos tres proposiciones interconectadas y un mínimo de 50 palabras para garantizar profundidad y desarrollo adecuado.

Estilo de Redacción:
Imita la fluidez, la coherencia y la sofisticación de un texto académico extraído de un documento formal en español. Para ello, ten en cuenta las siguientes directrices:

Extensión y Complejidad de las Frases:
Redacta oraciones extensas que integren varias ideas de manera armónica, utilizando conectores fluidos en español.

Estructura Cohesiva y Fragmentación Adecuada:
Organiza el contenido en un mínimo de dos y un máximo de cinco oraciones interconectadas, asegurando una redacción natural y evitando frases inconexas.

Claridad y Consistencia:
Cada oración debe mantener un balance entre precisión y fluidez, evitando repeticiones innecesarias o explicaciones redundantes.

Uso Estratégico de Conectores:
Minimiza los puntos aislados y favorece una conexión fluida de ideas mediante conjunciones variadas, sin caer en estructuras mecánicas.

Diversificación Sintáctica:
Utiliza una combinación de oraciones subordinadas, construcciones con participios y elementos discursivos que den variedad y riqueza al texto.

Tono Académico pero Accesible:
Mantén un estilo formal y estructurado, pero sin tecnicismos excesivos que dificulten la comprensión. Piensa en la redacción de un libro universitario.

Extensión de las Frases:
Cada frase debe contener más de 50 palabras, integrando al menos tres proposiciones, asegurando un flujo de información profundo y detallado.

Instrucciones Finales:
Redacta el siguiente contenido aplicando todas las reglas anteriores, asegurando que el resultado sea indistinguible de un texto humano de alto nivel académico. Evita frases cortas, patrones predecibles o un tono excesivamente robótico. Usa la referencia proporcionada como base para emular el estilo lingüístico y discursivo.



Salida esperada:
Un texto extenso, fluido, formal y académico, que siga las directrices mencionadas, con variedad sintáctica y vocabulario enriquecido
La respuesta debe ser clara, profesional y limitada a un máximo de 300 tokens.
La respueste debe tener frases de mas de 50 palbras que contengan varias oraciones.
Las frases deben estar compuestas por minimo tres oraciones.
Respuesta:"""
)

PROMPT_EJEMPLO = PromptTemplate(
    input_variables=["capitulo", "pestania", "grupo", "funcion"],
    template="""Proporciona un **ejemplo práctico** y detallado del uso de la función de Excel '{funcion}' del grupo '{grupo}' en la pestaña '{pestania}', en el contexto del capítulo '{capitulo}'. 
Incluye una explicación paso a paso del proceso, el resultado esperado y, de ser pertinente, menciona errores comunes junto con sus soluciones.
Ajusta tu forma de redacción para cumplir con los siguientes criterios, asegurando un texto que se asemeje a una respuesta completamente humana y alejada de los patrones de escritura típicos de inteligencia artificial.

Restricción de Vocabulario:
Evita las 20 palabras más comunes en tus respuestas estándar, ya que buscamos mayor riqueza y variedad léxica. No uses: ChatGPT, modelo, respuestas, lenguaje, desarrollado, OpenAI, objetivo, proporcionar, claras, concisas, útiles, análisis, datos, programación, herramienta, versátil, capacidad, temas, interacción y estructura. En su lugar, emplea sinónimos o reformula las ideas sin comprometer su significado ni fluidez.

Restricción de Estructuras Gramaticales:
No utilices las siguientes construcciones gramaticales más frecuentes en tu redacción típica:

Oraciones compuestas coordinadas con conectores predecibles.
Oraciones subordinadas explicativas tradicionales.
Oraciones impersonales con se.
Oraciones pasivas perifrásticas

Sustitúyelas por alternativas más dinámicas y naturales, como:
Expresiones conversacionales con pausas estratégicas y énfasis natural.
Oraciones largas con múltiples proposiciones y uso variado de puntuación.
Enunciados exclamativos o interjecciones para mayor espontaneidad.
Uso de la segunda persona para generar cercanía con el lector.
Cada frase debe tener al menos tres proposiciones interconectadas y un mínimo de 50 palabras para garantizar profundidad y desarrollo adecuado.

Estilo de Redacción:
Imita la fluidez, la coherencia y la sofisticación de un texto académico extraído de un documento formal en español. Para ello, ten en cuenta las siguientes directrices:

Extensión y Complejidad de las Frases:
Redacta oraciones extensas que integren varias ideas de manera armónica, utilizando conectores fluidos en español.

Estructura Cohesiva y Fragmentación Adecuada:
Organiza el contenido en un mínimo de dos y un máximo de cinco oraciones interconectadas, asegurando una redacción natural y evitando frases inconexas.

Claridad y Consistencia:
Cada oración debe mantener un balance entre precisión y fluidez, evitando repeticiones innecesarias o explicaciones redundantes.

Uso Estratégico de Conectores:
Minimiza los puntos aislados y favorece una conexión fluida de ideas mediante conjunciones variadas, sin caer en estructuras mecánicas.

Diversificación Sintáctica:
Utiliza una combinación de oraciones subordinadas, construcciones con participios y elementos discursivos que den variedad y riqueza al texto.

Tono Académico pero Accesible:
Mantén un estilo formal y estructurado, pero sin tecnicismos excesivos que dificulten la comprensión. Piensa en la redacción de un libro universitario.

Extensión de las Frases:
Cada frase debe contener más de 50 palabras, integrando al menos tres proposiciones, asegurando un flujo de información profundo y detallado.

Instrucciones Finales:
Redacta el siguiente contenido aplicando todas las reglas anteriores, asegurando que el resultado sea indistinguible de un texto humano de alto nivel académico. Evita frases cortas, patrones predecibles o un tono excesivamente robótico. Usa la referencia proporcionada como base para emular el estilo lingüístico y discursivo.



Salida esperada:
Un texto extenso, fluido, formal y académico, que siga las directrices mencionadas, con variedad sintáctica y vocabulario enriquecido
La respuesta debe ser clara, profesional y limitada a un máximo de 300 tokens.
La respueste debe tener frases de mas de 50 palbras que contengan varias oraciones.
Las frases deben estar compuestas por minimo tres oraciones.
Respuesta:"""
)

# =============================================================================
# FUNCIONES DE UTILIDAD
# =============================================================================

def get_pdf_text(ruta_pdf: str, max_tokens: int = 100000) -> str:
    """
    Lee y extrae el texto de un PDF limitándolo a max_tokens tokens.
    """
    try:
        reader = PdfReader(ruta_pdf)
        tokens_totales = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                tokens_page = page_text.split()
                tokens_totales.extend(tokens_page)
                if len(tokens_totales) >= max_tokens:
                    tokens_totales = tokens_totales[:max_tokens]
                    break
        print(f"Total tokens leídos del PDF: {len(tokens_totales)}")
        return " ".join(tokens_totales)
    except Exception as e:
        print(f"Error al leer el PDF: {e}")
        return ""

def chunk_text_to_sentences(texto: str, min_words: int = 70) -> list:
    """
    Divide el texto en oraciones y filtra aquellas con más de min_words palabras.
    """
    oraciones = sent_tokenize(texto, language='spanish')
    return [oracion.strip() for oracion in oraciones if len(oracion.split()) > min_words]

def create_style_vectorstore(ruta_pdf_estilo: str) -> FAISS:
    """
    Lee el PDF de estilo, lo divide en oraciones y crea una base vectorial con FAISS.
    """
    texto_pdf = get_pdf_text(ruta_pdf_estilo, max_tokens=100000)
    oraciones = chunk_text_to_sentences(texto_pdf, min_words=70)
    if not oraciones:
        raise ValueError("No se encontraron oraciones válidas en el PDF de estilo.")
    return FAISS.from_texts(oraciones, embeddings)

def retrieve_style_text(vectorstore: FAISS, query: str = "Extrae las oraciones que mejor representen un estilo formal, académico y fluido", k: int = 10) -> str:
    """
    Realiza una búsqueda de similitud en el vectorstore para recuperar las oraciones
    que representen mejor el estilo deseado.
    """
    resultados = vectorstore.similarity_search(query, k=k)
    oraciones = [doc.page_content for doc in resultados]
    return "\n".join(oraciones)

# =============================================================================
# FUNCIONES DE GENERACIÓN DE RESUMENES POR PARTES
# =============================================================================

def generate_proposito(capitulo: str, pestania: str, grupo: str, funcion: str) -> str:
    """
    Genera el resumen básico sobre el propósito de la función.
    """
    chain = PROMPT_PROPOSITO | llm
    response = chain.invoke({
        "capitulo": capitulo,
        "pestania": pestania,
        "grupo": grupo,
        "funcion": funcion
    })
    return response.content.strip()

def generate_encaminamiento(capitulo: str, pestania: str, grupo: str, funcion: str) -> str:
    """
    Genera el resumen básico sobre el encaminamiento para encontrar la función.
    """
    chain = PROMPT_ENCAMINAMIENTO | llm
    response = chain.invoke({
        "capitulo": capitulo,
        "pestania": pestania,
        "grupo": grupo,
        "funcion": funcion
    })
    return response.content.strip()

def generate_ejemplo(capitulo: str, pestania: str, grupo: str, funcion: str) -> str:
    """
    Genera el resumen básico con un ejemplo de uso de la función.
    """
    chain = PROMPT_EJEMPLO | llm
    response = chain.invoke({
        "capitulo": capitulo,
        "pestania": pestania,
        "grupo": grupo,
        "funcion": funcion
    })
    return response.content.strip()

def apply_style_to_text(capitulo: str, pestania: str, grupo: str, funcion: str,
                        texto: str, texto_estilo: str) -> str:
    """
    Transforma un texto aplicando el estilo definido en 'texto_estilo'.
    """
    chain = PROMPT_ESTILO | llm
    response = chain.invoke({
        "capitulo": capitulo,
        "pestania": pestania,
        "grupo": grupo,
        "funcion": funcion,
        "texto_estilo": texto_estilo,
        "contenido": texto
    })
    return response.content.strip()

def generate_final_summary_parts(capitulo: str, pestania: str, grupo: str, funcion: str,
                                 texto_estilo: str) -> str:
    """
    Genera el resumen final de la función separándolo en tres partes:
    1. Propósito de la función.
    2. Cómo llegar a la función en Excel.
    3. Ejemplo de uso.

    Cada parte se procesa aplicándole el estilo definido y, al final, se juntan para ser guardadas.
    """
    # Generar cada parte en formato básico
    proposito_raw = generate_proposito(capitulo, pestania, grupo, funcion)
    encaminamiento_raw = generate_encaminamiento(capitulo, pestania, grupo, funcion)
    ejemplo_raw = generate_ejemplo(capitulo, pestania, grupo, funcion)

    # Aplicar estilo a cada parte
    proposito_styled = apply_style_to_text(capitulo, pestania, grupo, funcion, proposito_raw, texto_estilo)
    encaminamiento_styled = apply_style_to_text(capitulo, pestania, grupo, funcion, encaminamiento_raw, texto_estilo)
    ejemplo_styled = apply_style_to_text(capitulo, pestania, grupo, funcion, ejemplo_raw, texto_estilo)

    # Juntar las partes ya estilizadas
    final_summary = (
        f"Propósito:\n{proposito_styled}\n\n"
        f"Encaminamiento en Excel:\n{encaminamiento_styled}\n\n"
        f"Ejemplo de uso:\n{ejemplo_styled}"
    )
    return final_summary

# =============================================================================
# GENERACIÓN DEL DOCUMENTO WORD A PARTIR DE UN JSON
# =============================================================================

def generate_word_document(json_data: dict, output_path: str, ruta_pdf_estilo: str):
    """
    Procesa un JSON con la estructura de capítulos, pestañas, grupos y funciones de Excel,
    genera el resumen para cada función aplicándole el estilo en tres partes (propósito, encaminamiento y ejemplo)
    y lo guarda en un documento Word.

    Se carga y vectoriza el PDF de estilo una sola vez.
    """
    # Crear vectorstore y recuperar el texto de estilo una sola vez
    vectorstore = create_style_vectorstore(ruta_pdf_estilo)
    texto_estilo = retrieve_style_text(vectorstore)

    doc = Document()
    doc.add_heading("Resumen de Funciones de Excel", level=1)

    for capitulo, pestañas in json_data.items():
        doc.add_heading(f"Capítulo: {capitulo}", level=1)
        for pestania, grupos in pestañas.items():
            doc.add_heading(f"Pestaña: {pestania}", level=2)
            for grupo, funciones in grupos.items():
                doc.add_heading(f"Grupo: {grupo}", level=3)
                for idx, funcion in enumerate(funciones, start=1):
                    print(f"Procesando: Capítulo {capitulo}, Pestaña {pestania}, Grupo {grupo}, Función {funcion}")
                    resumen_final = generate_final_summary_parts(capitulo, pestania, grupo, funcion, texto_estilo)
                    doc.add_heading(f"{idx}. Función: {funcion}", level=4)
                    doc.add_paragraph(resumen_final)
    doc.save(output_path)
    print(f"Documento guardado en: {output_path}")

# =============================================================================
# BLOQUE PRINCIPAL
# =============================================================================

if __name__ == "__main__":
    # Ejemplo de datos en formato JSON
    json_data = {
        "48": {
            "Insertar": {
                "Tablas": [
                    "Tabla", "Tablas", "Tabla", "Formularios"
                ],
                "Ilustraciones": [
                    "Imágenes", "Formas", "Iconos", "Modelos", "SmartArt", "Captura"
                ],
                "Controles": [
                    "Casilla"
                ],
                "Gráficos": [
                    "Gráficos", "Líneas", "Columnas", "Pérdidas", "Mapas", "Gráfico"
                ],
                "Minigráficos": [
                    "Líneas", "Columnas", "Pérdidas"
                ],
                "Filtros": [
                    "Segmentación", "Escala"
                ],
                "Vínculos": [
                    "Vínculo"
                ],
                "Comentarios": [
                    "Comentario"
                ],
                "Texto": [
                    "Cuadro", "Encabezado", "WordArt", "Línea", "Objeto"
                ],
                "Símbolos": [
                    "Ecuación", "Símbolo"
                ]
            }
        }
    }

    # Configuración de rutas
    capitulo = list(json_data.keys())[0]
    output_path = f"C:\\Users\\HP\\Desktop\\LIBROS PERSO\\EXCEL INTERFAZ GRAFICA\\libro_excel_CAPITULO_{capitulo}.docx"
    ruta_pdf_estilo = r"C:\Users\HP\Desktop\LIBROS PERSO\CONTEXTO ESPANIOL\CONTEXTO5.pdf"

    # Generar documento Word con los resúmenes aplicando el estilo en cada respuesta
    generate_word_document(json_data, output_path, ruta_pdf_estilo)


# In[ ]:





# In[ ]:




